import io

import docx

from app.ingestion.extractors.docx import extract


def _make_docx_with_body_and_table(body_text: str, table_cells: list[list[str]]) -> bytes:
    document = docx.Document()
    document.add_paragraph(body_text)

    if table_cells:
        rows = len(table_cells)
        cols = len(table_cells[0])
        table = document.add_table(rows=rows, cols=cols)
        for row_index, row_values in enumerate(table_cells):
            for col_index, value in enumerate(row_values):
                table.cell(row_index, col_index).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_body_text_extracted_as_text_chunks():
    docx_bytes = _make_docx_with_body_and_table(
        "This is the main body paragraph.", table_cells=[]
    )

    result = extract(docx_bytes, "sample.docx")

    assert any(c["text"] == "This is the main body paragraph." for c in result)
    assert all(c["extraction_method"] == "text" for c in result)


def test_table_cell_text_is_not_dropped():
    docx_bytes = _make_docx_with_body_and_table(
        "Body paragraph.",
        table_cells=[["Revenue", "1000"], ["Expenses", "500"]],
    )

    result = extract(docx_bytes, "sample.docx")
    texts = [c["text"] for c in result]

    assert "Revenue" in texts
    assert "1000" in texts
    assert "Expenses" in texts
    assert "500" in texts


def test_page_is_none_and_location_is_populated():
    docx_bytes = _make_docx_with_body_and_table(
        "Body paragraph.", table_cells=[["A", "B"]]
    )

    result = extract(docx_bytes, "sample.docx")

    for chunk in result:
        assert chunk["page"] is None
        assert chunk["location"] is not None


def test_empty_paragraphs_are_skipped():
    document = docx.Document()
    document.add_paragraph("Real content here.")
    document.add_paragraph("")  # empty paragraph, should be skipped
    buffer = io.BytesIO()
    document.save(buffer)

    result = extract(buffer.getvalue(), "sample.docx")

    assert len(result) == 1
    assert result[0]["text"] == "Real content here."