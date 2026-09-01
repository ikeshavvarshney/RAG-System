import io

import docx


def extract(content: bytes, filename: str):
    """Extract body text and table-cell text from a DOCX file.

    DOCX has no real page concept, so `page` is always None; `location`
    carries a positional hint instead (paragraph index or table position).
    """
    document = docx.Document(io.BytesIO(content))
    chunks = []

    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if text:
            chunks.append({
                "page": None,
                "location": f"paragraph_{index}",
                "text": text,
                "extraction_method": "text",
            })

    # Table-cell text is NOT included in document.paragraphs — it must be
    # walked separately via document.tables, or it silently vanishes.
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    chunks.append({
                        "page": None,
                        "location": f"table_{table_index}_row_{row_index}_cell_{cell_index}",
                        "text": text,
                        "extraction_method": "text",
                    })

    return chunks